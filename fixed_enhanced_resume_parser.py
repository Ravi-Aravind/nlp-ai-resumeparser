
import re
import os
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
import asyncio

# NLP imports with error handling
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    logging.warning("spaCy not available. Install with: pip install spacy")
    SPACY_AVAILABLE = False

# PDF parsing imports
try:
    import PyPDF2
    PDF_PARSING_AVAILABLE = True
except ImportError:
    try:
        import pypdf  # Alternative PDF library
        import pypdf as PyPDF2
        PDF_PARSING_AVAILABLE = True
    except ImportError:
        logging.warning("PDF parsing not available. Install with: pip install PyPDF2")
        PDF_PARSING_AVAILABLE = False

# DOCX parsing imports
try:
    from docx import Document
    DOCX_PARSING_AVAILABLE = True
except ImportError:
    logging.warning("DOCX parsing not available. Install with: pip install python-docx")
    DOCX_PARSING_AVAILABLE = False

logger = logging.getLogger(__name__)

class EnhancedResumeParser:
    """Enhanced resume parser with robust error handling and fallbacks"""

    def __init__(self):
        self.nlp = None
        self._initialize_nlp()

        # Enhanced skill taxonomy
        self.skill_taxonomy = {
            "programming_languages": [
                "Python", "Java", "JavaScript", "C++", "C#", "Ruby", "PHP", "Go", "Rust",
                "TypeScript", "Swift", "Kotlin", "Scala", "R", "MATLAB", "Perl", "C",
                "Objective-C", "Dart", "Julia", "F#", "Haskell"
            ],
            "web_technologies": [
                "HTML", "CSS", "React", "Angular", "Vue.js", "Node.js", "Express.js",
                "Django", "Flask", "Spring Boot", "ASP.NET", "Laravel", "Ruby on Rails",
                "Next.js", "Nuxt.js", "Svelte", "Bootstrap", "Tailwind CSS"
            ],
            "databases": [
                "MySQL", "PostgreSQL", "MongoDB", "Redis", "Cassandra", "DynamoDB",
                "Oracle", "SQL Server", "SQLite", "Elasticsearch", "MariaDB", 
                "CouchDB", "InfluxDB"
            ],
            "cloud_platforms": [
                "AWS", "Azure", "Google Cloud", "GCP", "Heroku", "DigitalOcean", 
                "Linode", "Vercel", "Netlify", "Firebase"
            ],
            "devops_tools": [
                "Docker", "Kubernetes", "Jenkins", "GitLab CI", "Travis CI", "CircleCI",
                "Terraform", "Ansible", "Puppet", "Chef", "Vagrant", "Prometheus", 
                "Grafana"
            ],
            "version_control": [
                "Git", "GitHub", "GitLab", "Bitbucket", "SVN", "Mercurial"
            ],
            "soft_skills": [
                "Communication", "Leadership", "Problem Solving", "Teamwork", "Agile",
                "Scrum", "Project Management", "Time Management", "Critical Thinking",
                "Collaboration", "Adaptability"
            ],
            "frameworks": [
                "TensorFlow", "PyTorch", "Keras", "Pandas", "NumPy", "Scikit-learn",
                "OpenCV", "Matplotlib", "Seaborn", "Plotly"
            ]
        }

        # Flatten skills for easier matching
        self.all_skills = []
        for category in self.skill_taxonomy.values():
            self.all_skills.extend(category)

    def _initialize_nlp(self):
        """Initialize NLP model with fallback"""
        if SPACY_AVAILABLE:
            try:
                self.nlp = spacy.load("en_core_web_sm")
                logger.info("spaCy model loaded successfully")
            except OSError:
                logger.warning("spaCy model not found. Attempting to download...")
                try:
                    os.system("python -m spacy download en_core_web_sm")
                    self.nlp = spacy.load("en_core_web_sm")
                    logger.info("spaCy model downloaded and loaded")
                except Exception as e:
                    logger.error(f"Failed to download spaCy model: {e}")
                    self.nlp = None
        else:
            self.nlp = None

    async def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Parse resume from file path with comprehensive error handling"""
        try:
            # Extract text based on file type
            text = await self._extract_text_from_file(file_path)

            if not text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return self._create_empty_result()

            # Parse different sections
            parsed_data = {
                "name": self._extract_name(text),
                "email": self._extract_email(text),
                "phone": self._extract_phone(text),
                "skills": self._extract_skills(text),
                "experience": self._extract_experience(text),
                "education": self._extract_education(text),
                "work_history": self._extract_work_history(text),
                "location": self._extract_location(text),
                "salary_expectation": self._extract_salary_expectation(text),
                "raw_text": text[:2000],  # Limit raw text size
                "confidence_scores": {}
            }

            # Calculate confidence scores
            parsed_data["confidence_scores"] = self._calculate_confidence_scores(parsed_data, text)

            logger.info(f"Resume parsed successfully: {parsed_data.get('name', 'Unknown')}")
            return parsed_data

        except Exception as e:
            logger.error(f"Resume parsing error: {e}")
            return self._create_empty_result()

    async def _extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file formats"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()

            if file_extension == '.pdf':
                return await self._extract_pdf_text(file_path)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_docx_text(file_path)
            elif file_extension == '.txt':
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return await f.read()
            else:
                # Try to read as plain text
                try:
                    async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return await f.read()
                except:
                    logger.error(f"Unsupported file format: {file_extension}")
                    return ""

        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return ""

    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        if not PDF_PARSING_AVAILABLE:
            return "PDF parsing not available. Please install PyPDF2."

        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            return text.strip()

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""

    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCX_PARSING_AVAILABLE:
            return "DOCX parsing not available. Please install python-docx."

        try:
            doc = Document(file_path)
            text_parts = []

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""

    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name with improved logic"""
        lines = text.split('\n')[:10]  # Check first 10 lines

        # Enhanced name patterns
        name_patterns = [
            r'^\s*([A-Z][a-z]+(?:\s+[A-Z][a-z']+){1,2})\s*$',  # Full name on its own line
            r'Name:\s*([A-Z][a-z]+(?:\s+[A-Z][a-z']+){0,2})',    # "Name: John Doe"
            r'^\s*([A-Z][a-z]+(?:\s+[A-Z]\.?\s*[A-Z][a-z]+))\s*$',  # Name with middle initial
        ]

        for line in lines:
            line = line.strip()
            if len(line) < 3 or len(line) > 50:
                continue

            # Skip headers and common resume words
            skip_words = ['resume', 'cv', 'curriculum', 'vitae', 'professional', 'summary', 'objective']
            if any(word in line.lower() for word in skip_words):
                continue

            # Try each pattern
            for pattern in name_patterns:
                match = re.search(pattern, line)
                if match:
                    potential_name = match.group(1).strip()

                    # Additional validation
                    words = potential_name.split()
                    if 1 <= len(words) <= 3:
                        # Check if it looks like a real name
                        if all(word[0].isupper() and word[1:].islower() for word in words if len(word) > 1):
                            return potential_name

        return None

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)

        # Filter out common false positives
        valid_emails = []
        for email in matches:
            email_lower = email.lower()
            if not any(domain in email_lower for domain in ['example.com', 'test.com', 'placeholder']):
                valid_emails.append(email)

        return valid_emails[0].lower() if valid_emails else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number with multiple formats"""
        phone_patterns = [
            r'\+?\d{1,3}[-\s]?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{4}',  # International format
            r'\(?\d{3}\)?[-\.\s]?\d{3}[-\.\s]?\d{4}',                # US format
            r'Phone:?\s*([\d\s\-\(\)\+\.]+)',                         # Labeled phone
            r'Mobile:?\s*([\d\s\-\(\)\+\.]+)',                        # Mobile label
            r'Tel:?\s*([\d\s\-\(\)\+\.]+)'                           # Tel label
        ]

        for pattern in phone_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # Clean up the phone number
                phone = re.sub(r'[^\d\+]', '', match)
                if 10 <= len(phone) <= 15:  # Valid phone number length
                    return match.strip()

        return None

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills with improved matching"""
        found_skills = set()
        text_lower = text.lower()

        # Direct skill matching
        for skill in self.all_skills:
            skill_lower = skill.lower()
            # Use word boundaries for better matching
            pattern = r'\b' + re.escape(skill_lower) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill)

        # Extract from skills section specifically
        skills_section_match = re.search(
            r'(?:technical\s+skills?|skills?|technologies?|competencies?)\s*:?\s*([\s\S]*?)(?=\n\s*\n|\n[A-Z]|$)',
            text, re.IGNORECASE
        )

        if skills_section_match:
            skills_text = skills_section_match.group(1)
            # Split by common delimiters
            skill_candidates = re.split(r'[,;|•\n\-]+', skills_text)

            for candidate in skill_candidates:
                candidate = candidate.strip()
                if candidate and len(candidate) <= 30:  # Reasonable skill name length
                    # Check against known skills
                    for skill in self.all_skills:
                        if skill.lower() in candidate.lower():
                            found_skills.add(skill)

        # Use NLP if available for additional extraction
        if self.nlp and SPACY_AVAILABLE:
            try:
                doc = self.nlp(text)
                for ent in doc.ents:
                    if ent.label_ in ['ORG', 'PRODUCT'] and len(ent.text) < 30:
                        # Check if entity matches known skills
                        for skill in self.all_skills:
                            if skill.lower() == ent.text.lower():
                                found_skills.add(skill)
            except Exception as e:
                logger.warning(f"NLP extraction error: {e}")

        return sorted(list(found_skills))

    def _extract_experience(self, text: str) -> Optional[str]:
        """Extract years of experience with improved patterns"""
        # Direct experience mentions
        exp_patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience:?\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s+(?:of\s+)?experience',
            r'over\s+(\d+)\s+years?',
            r'more\s+than\s+(\d+)\s+years?'
        ]

        for pattern in exp_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                years = int(match.group(1))
                return f"{years}+ years"

        # Calculate from work history dates
        work_years = self._calculate_experience_from_dates(text)
        if work_years:
            return work_years

        return None

    def _calculate_experience_from_dates(self, text: str) -> Optional[str]:
        """Calculate experience from work history dates"""
        try:
            # Find all year mentions (1990-2030)
            years = re.findall(r'\b(19\d{2}|20[0-3]\d)\b', text)

            if len(years) >= 2:
                year_ints = [int(y) for y in years]
                start_year = min(year_ints)

                # Use current year or latest year in resume
                current_year = datetime.now().year
                end_year = min(max(year_ints), current_year)

                experience_years = end_year - start_year

                if 0 < experience_years <= 50:  # Sanity check
                    return f"{experience_years}+ years"

            return None

        except Exception as e:
            logger.error(f"Date calculation error: {e}")
            return None

    def _extract_education(self, text: str) -> Optional[str]:
        """Extract education information"""
        education_patterns = [
            r'(Bachelor[\s\']?s?(?:\s+of\s+[\w\s]+)?|B\.?[AS]\.?(?:\s+[\w\s]+)?)',
            r'(Master[\s\']?s?(?:\s+of\s+[\w\s]+)?|M\.?[AS]\.?(?:\s+[\w\s]+)?)',
            r'(PhD|Ph\.?D\.?|Doctor(?:ate)?(?:\s+of\s+[\w\s]+)?)',
            r'(Associate[\s\']?s?(?:\s+[\w\s]+)?)',
            r'(Certificate(?:\s+[\w\s]+)?)',
            r'(Diploma(?:\s+[\w\s]+)?)'
        ]

        education_info = []

        for pattern in education_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            education_info.extend(matches)

        # Look for university names
        university_pattern = r'(\b(?:University|College|Institute|School)(?:\s+of\s+[\w\s]+)?)\b'
        universities = re.findall(university_pattern, text, re.IGNORECASE)
        education_info.extend(universities[:2])  # Limit to first 2

        if education_info:
            return "; ".join(education_info[:3])  # Limit to first 3 items

        return None

    def _extract_work_history(self, text: str) -> List[Dict[str, Any]]:
        """Extract work history with improved parsing"""
        work_history = []

        # Find experience/work section
        experience_section = re.search(
            r'(?:professional\s+experience|work\s+experience|experience|employment)\s*:?\s*([\s\S]*?)(?=\n\s*(?:education|skills|projects|certifications)|$)',
            text, re.IGNORECASE
        )

        if not experience_section:
            return work_history

        experience_text = experience_section.group(1)
        lines = experience_text.split('\n')

        current_job = {}
        description_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Job title and company patterns
            job_patterns = [
                r'^([^|]+)\|\s*([^|]+)\|?\s*(\d{4}\s*-\s*\d{4}|\d{4}\s*-\s*Present)?',  # Title | Company | Dates
                r'^([^,]+),\s*([^,]+)\s*(\d{4}\s*-\s*\d{4}|\d{4}\s*-\s*Present)?',        # Title, Company Dates
                r'^((?:[A-Z][a-z]*\s*)+)\s*(?:at|@)\s*([^\d]*)(\d{4}\s*-\s*\d{4}|\d{4}\s*-\s*Present)?'  # Title at Company Dates
            ]

            job_matched = False
            for pattern in job_patterns:
                match = re.search(pattern, line)
                if match:
                    # Save previous job if exists
                    if current_job:
                        if description_lines:
                            current_job['description'] = description_lines[:3]  # Limit description
                        work_history.append(current_job)
                        description_lines = []

                    # Start new job
                    current_job = {
                        'title': match.group(1).strip(),
                        'company': match.group(2).strip() if len(match.groups()) > 1 else '',
                        'dates': match.group(3).strip() if len(match.groups()) > 2 and match.group(3) else ''
                    }
                    job_matched = True
                    break

            # Collect description lines
            if not job_matched and current_job and len(description_lines) < 5:
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    description_lines.append(line[1:].strip())
                elif len(line) > 20:  # Substantial description line
                    description_lines.append(line)

        # Save the last job
        if current_job:
            if description_lines:
                current_job['description'] = description_lines[:3]
            work_history.append(current_job)

        return work_history[:5]  # Limit to 5 most recent positions

    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location information"""
        location_patterns = [
            r'Location:?\s*([A-Za-z\s,]+(?:,\s*[A-Z]{2})?)',
            r'Address:?\s*([A-Za-z\s,]+(?:,\s*[A-Z]{2})?)',
            r'([A-Za-z\s]+,\s*[A-Z]{2})(?:\s|$)',  # City, State
            r'([A-Za-z\s]+,\s*(?:California|New York|Texas|Florida|India|Canada|UK|USA))\b'
        ]

        for pattern in location_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                location = match.group(1).strip()
                # Clean up location
                location = re.sub(r'\s+', ' ', location)
                if len(location) > 3 and len(location) < 100:
                    return location

        return None

    def _extract_salary_expectation(self, text: str) -> Optional[str]:
        """Extract salary expectation"""
        salary_patterns = [
            r'Salary:?\s*\$?([\d,]+(?:\.\d+)?(?:k|K)?)',
            r'Expected\s+Salary:?\s*\$?([\d,]+(?:\.\d+)?(?:k|K)?)',
            r'Compensation:?\s*\$?([\d,]+(?:\.\d+)?(?:k|K)?)',
            r'\$([\d,]+(?:\.\d+)?(?:k|K)?)\s*(?:per\s+year|annually|yearly)?'
        ]

        for pattern in salary_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)

        return None

    def _calculate_confidence_scores(self, parsed_data: Dict[str, Any], text: str) -> Dict[str, float]:
        """Calculate confidence scores for extraction accuracy"""
        scores = {}

        # Name confidence
        name = parsed_data.get('name')
        if name and len(name.split()) >= 2:
            scores['name'] = 0.9
        elif name:
            scores['name'] = 0.6
        else:
            scores['name'] = 0.1

        # Email confidence
        scores['email'] = 0.95 if parsed_data.get('email') else 0.0

        # Phone confidence
        scores['phone'] = 0.85 if parsed_data.get('phone') else 0.0

        # Skills confidence
        skills = parsed_data.get('skills', [])
        if len(skills) >= 5:
            scores['skills'] = 0.9
        elif len(skills) >= 3:
            scores['skills'] = 0.7
        elif len(skills) >= 1:
            scores['skills'] = 0.5
        else:
            scores['skills'] = 0.1

        # Experience confidence
        scores['experience'] = 0.8 if parsed_data.get('experience') else 0.3

        # Education confidence
        scores['education'] = 0.8 if parsed_data.get('education') else 0.2

        # Overall confidence
        scores['overall'] = sum(scores.values()) / len(scores)

        return scores

    def _create_empty_result(self) -> Dict[str, Any]:
        """Create empty result structure for failed parsing"""
        return {
            "name": None,
            "email": None,
            "phone": None,
            "skills": [],
            "experience": None,
            "education": None,
            "work_history": [],
            "location": None,
            "salary_expectation": None,
            "raw_text": "",
            "confidence_scores": {
                "name": 0.0,
                "email": 0.0,
                "phone": 0.0,
                "skills": 0.0,
                "experience": 0.0,
                "education": 0.0,
                "overall": 0.0
            }
        }

# Example usage and testing
async def test_parser():
    """Test the resume parser"""
    parser = EnhancedResumeParser()

    # Test with the provided resume.docx if available
    test_files = ['resume.docx', 'sample_resume.pdf', 'test_resume.txt']

    for test_file in test_files:
        if os.path.exists(test_file):
            try:
                result = await parser.parse_resume(test_file)
                print(f"\n✓ Parsed {test_file}:")
                print(f"  Name: {result.get('name')}")
                print(f"  Email: {result.get('email')}")
                print(f"  Skills: {len(result.get('skills', []))} found")
                print(f"  Confidence: {result.get('confidence_scores', {}).get('overall', 0):.2f}")
            except Exception as e:
                print(f"✗ Error parsing {test_file}: {e}")

if __name__ == "__main__":
    # Run test if executed directly
    asyncio.run(test_parser())
